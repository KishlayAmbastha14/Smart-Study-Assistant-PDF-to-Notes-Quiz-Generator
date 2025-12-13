import streamlit as st
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_groq import ChatGroq
from langchain.prompts import PromptTemplate
from langchain_core.runnables import RunnableParallel
from langchain_core.output_parsers import StrOutputParser
from dotenv import load_dotenv
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import CharacterTextSplitter
from langchain_text_splitters import RecursiveCharacterTextSplitter
import tempfile
import os
from docx import Document
from io import BytesIO


load_dotenv()
groq_api_key = os.getenv("GROQ_API_KEY")
os.environ["GOOGLE_API_KEY"] = os.getenv("GOOGLE_API_KEY")

model1 = ChatGoogleGenerativeAI(model="gemini-2.5-flash")
# model2 = ChatGroq(model="gemma2-9b-it",groq_api_key=groq_api_key)
model2 = ChatGoogleGenerativeAI(model="gemini-2.5-pro")


summary_length = st.radio("Choose summary length:", ["Short","Medium","Detailed"])

length_instruction = {
    "Short": "Write a 3-5 sentence summary.",
    "Medium": "Write one detailed paragraph.",
    "Detailed": "Write a detailed summary around one page."
}[summary_length]

prompt1 = PromptTemplate(
    template=f"Generate notes based on the text.\n{length_instruction}\nText: {{text}}",
    input_variables=["text"]
)

# prompt1 = PromptTemplate(
#   template="Generate short and simple notes from the following text \n {text}",
#   input_variables=["text"]
# )

prompt2 = PromptTemplate(
  template="Generate 5 questions answers from the following text \n {text}",
  input_variables=["text"]
)

prompt3 = PromptTemplate(
  template="Merge the provided notes and quiz into a single document \n notes -> {notes} and quiz -> {quiz}",
  input_variables=["notes","quiz"]
)

parser = StrOutputParser()

parallel_chain = RunnableParallel({
  'notes' : prompt1 | model1 | parser,
  'quiz' : prompt2 | model2 | parser
})

merge_chain = prompt3 | model1 | parser

chain = parallel_chain | merge_chain


# --------------- STREAMLIT SET_UP ---------------------
st.set_page_config(page_title="Smart Study Assistant", layout="wide")
st.title("📘 Smart Study Assistant")

# --- user_choice ----
input_mode = st.radio("Chooose the input:",("Text","PDF"))
final_text = None

# ---- user_Choice is text -----
if input_mode == "Text":
  user_text = st.text_area("✍️ Enter your text here to generate notes and quiz:", height=200)
  if user_text.strip():
    final_text = user_text

elif input_mode == "PDF":
  uploaded_pdf = st.file_uploader("Upload a PDF", type=["pdf"])
  if uploaded_pdf is not None:
    with tempfile.NamedTemporaryFile(delete=False,suffix=".pdf") as temp_file: 
      temp_file.write(uploaded_pdf.read()) 
      temp_file_path = temp_file.name 
      loader = PyPDFLoader(temp_file_path)
      document = loader.load()
      text_splitter = RecursiveCharacterTextSplitter(chunk_size=2500,chunk_overlap =80,separators=["\n\n", "\n", ".", " "])
      split_chunks = text_splitter.split_documents(document)
    # st.write(len(split_chunks))
      final_text = " ".join([chunk.page_content for chunk in split_chunks])
    # st.write(len(final_text))


if st.button("Generate Study Guide"):
  if final_text:
    with st.spinner("Generating your study guide..."):
    # st.write(len(final_text))
      result = parallel_chain.invoke({"text":final_text})
      notes = result['notes']
      quiz = result['quiz']

      st.subheader("📘 Short Notes")
      st.markdown(notes)

      st.subheader("❓ Quiz (5 Questions)")
      st.markdown(quiz)

      merged_all = f"📘 Short Notes\n\n{notes}\n\n❓ Quiz\n\n{quiz}"

      # --- you can give quiz also ---

    #   st.subheader("📝 Take the quiz")
    #   user_answers = []
    #   questions = quiz.split("\n")
    #   st.write("Please answer the following questions:")
    # # st.write(questions)
    #   for q in questions:
    #       if q.strip():
    #         answer = st.text_input(q,key=q)
    #         user_answers.append(answer)
    #       ## user_answers.append((q,answer))
      
    #   if st.button("Submit Quiz"):
    #     st.write("✅ Correct Answers:")
    #     st.markdown(quiz)  # show original quiz with answers


    # --- DOWNLOAD OPTION ---
      doc = Document()
      doc.add_heading("Smart Study Guide",level=1)
      doc.add_paragraph(merged_all)

      buffer = BytesIO()
      doc.save(buffer)
      buffer.seek(0)

      st.download_button(
            label="📥 Download Notes as Word",
            data=buffer,
            file_name="study_guide.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

      st.subheader("📑 Generated Study Guide")
    # st.write(result)
    
    # st.text_area("Result", value=result, height=300)

else:
  st.warning("⚠️ Please enter text or upload the pdf before generating.")
